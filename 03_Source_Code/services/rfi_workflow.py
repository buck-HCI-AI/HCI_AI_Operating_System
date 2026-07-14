"""
HCI AI - RFI Workflow Service

End-to-end RFI handling: generate the Hendrickson Construction RFI Word
document (matching the real template found in each project's own "RFIs"
Drive folder), save it to the correct project location, update that
project's real RFI Log tracker tab, and create an Outlook draft with the
document attached.

Built 2026-07-09/10 per Chief Architect direction: the backend previously
only supported creating an RFI database row (see /field/rfi in
api/routers/gbt_gateway.py) - no update path, no document generation, no
Drive save, no tracker update, no email draft existed. Field GPT's report
that it could not do any of this was accurate, not a capability-exposure
bug (see memory: project_rfi_capability_gap_2026-07-10).
"""
import io, os, sys, json
from datetime import date, datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "integrations"))
sys.path.insert(0, os.path.dirname(__file__))

import psycopg2, psycopg2.extras
from dotenv import load_dotenv
from credentials import get_google_token
import requests

load_dotenv(os.path.join(os.path.dirname(__file__), "..", "..", ".env"))

DB = dict(
    host=os.environ.get("POSTGRES_HOST", "localhost"),
    port=int(os.environ.get("POSTGRES_PORT", 5432)),
    dbname=os.environ.get("POSTGRES_DB", "hci_os"),
    user=os.environ.get("POSTGRES_USER", "hci_admin"),
    password=os.environ.get("POSTGRES_PASSWORD", ""),
)


def _pg():
    return psycopg2.connect(**DB, cursor_factory=psycopg2.extras.RealDictCursor)


def get_rfi(rfi_id: int) -> dict:
    with _pg() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT r.*, p.project_code, p.name AS project_name
                FROM rfis r JOIN projects p ON p.id = r.project_id
                WHERE r.id = %s
            """, (rfi_id,))
            row = cur.fetchone()
            return dict(row) if row else None


def update_rfi(rfi_id: int, status: str = None, response: str = None,
                response_date: str = None, updated_by: str = "system",
                rfi_document_id: str = None, email_draft_id: str = None,
                response_document_id: str = None, originating_discipline: str = None,
                responsible_consultant: str = None, drawing_references: list = None,
                spec_references: list = None, cost_impact: float = None,
                schedule_impact_days: float = None, linked_change_order: str = None,
                linked_bulletin: str = None, linked_revision: str = None) -> dict:
    """
    The one real gap Field GPT correctly reported as missing: there was no
    way to update an existing RFI's status/response at all, only create new
    ones. Also logs the update to project_events for the same timeline
    every other project action feeds.

    Extended 2026-07-14 per GBT's RFI Tracker Refactor architecture: the
    tracker should be the master record, not just an Excel log, so every
    field GBT specified (discipline, consultant, drawing/spec refs, cost/
    schedule impact, linked change order/bulletin/revision, and the
    generated document + email draft IDs) can be set here rather than
    only existing as one-time API response evidence in run_rfi_workflow().
    """
    fields, values = [], []
    if status is not None:
        fields.append("status = %s"); values.append(status)
    if response is not None:
        fields.append("response = %s"); values.append(response)
    if response_date is not None:
        fields.append("response_date = %s"); values.append(response_date)
    if rfi_document_id is not None:
        fields.append("rfi_document_id = %s"); values.append(rfi_document_id)
    if email_draft_id is not None:
        fields.append("email_draft_id = %s"); values.append(email_draft_id)
    if response_document_id is not None:
        fields.append("response_document_id = %s"); values.append(response_document_id)
    if originating_discipline is not None:
        fields.append("originating_discipline = %s"); values.append(originating_discipline)
    if responsible_consultant is not None:
        fields.append("responsible_consultant = %s"); values.append(responsible_consultant)
    if drawing_references is not None:
        fields.append("drawing_references = %s"); values.append(drawing_references)
    if spec_references is not None:
        fields.append("spec_references = %s"); values.append(spec_references)
    if cost_impact is not None:
        fields.append("cost_impact = %s"); values.append(cost_impact)
    if schedule_impact_days is not None:
        fields.append("schedule_impact_days = %s"); values.append(schedule_impact_days)
    if linked_change_order is not None:
        fields.append("linked_change_order = %s"); values.append(linked_change_order)
    if linked_bulletin is not None:
        fields.append("linked_bulletin = %s"); values.append(linked_bulletin)
    if linked_revision is not None:
        fields.append("linked_revision = %s"); values.append(linked_revision)
    if not fields:
        return {"error": "nothing to update"}
    values.append(rfi_id)
    with _pg() as conn:
        conn.autocommit = True
        with conn.cursor() as cur:
            cur.execute(f"UPDATE rfis SET {', '.join(fields)} WHERE id = %s RETURNING *", values)
            row = cur.fetchone()
            if not row:
                return {"error": f"RFI {rfi_id} not found"}
            cur.execute("""
                INSERT INTO project_events (project_id, event_type, event_date, title, description, source_table, source_id, created_by)
                VALUES (%s, 'rfi_updated', CURRENT_DATE, %s, %s, 'rfis', %s, %s)
            """, (row["project_id"], f"RFI {row['rfi_number']} updated: {status or 'response added'}",
                  response or "", rfi_id, updated_by))
    return dict(row)


def generate_rfi_docx(rfi: dict) -> bytes:
    """
    Builds a Word document matching Hendrickson Construction's real RFI
    form, as found in each project's own Drive ("RFI Template" tab of the
    project's RFI Log workbook - see e.g. 1355 Riverside RFI - Log .xlsx,
    Drive id 1l47B4kQGMhfZLoJz6m3fy9EZFryGMyd4). Same layout/field labels,
    not an invented format.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)

    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("HENDRICKSON CONSTRUCTION INC.")
    run.bold = True
    run.font.size = Pt(14)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = h2.add_run("REQUEST FOR INFORMATION")
    run2.bold = True
    run2.font.size = Pt(12)
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].text, cells[1].text = "RFI Title:", rfi.get("subject", "")
    cells[2].text, cells[3].text = "RFI #:", rfi.get("rfi_number", "")
    r2 = table.rows[1].cells
    r2[0].text, r2[1].text = "Date:", str(rfi.get("submitted_date", ""))
    r2[2].text, r2[3].text = "Sub RFI#:", ""
    r3 = table.rows[2].cells
    r3[0].text, r3[1].text = "Concerned Party:", rfi.get("submitted_to", "")
    r3[2].text, r3[3].text = "Sent By:", rfi.get("submitted_by", "")
    r4 = table.rows[3].cells
    r4[0].text, r4[1].text = "Reference Page:", rfi.get("reference_page", "") or ""
    r4[2].text, r4[3].text = "Project:", f"{rfi.get('project_name', '')} ({rfi.get('project_code', '')})"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Hendrickson Construction is Requesting the Following:").bold = True
    doc.add_paragraph(rfi.get("question", ""))
    doc.add_paragraph()
    doc.add_paragraph()

    p2 = doc.add_paragraph()
    resp_run = p2.add_run("Response:")
    resp_run.bold = True
    doc.add_paragraph(rfi.get("response", "") or "[Awaiting response]")
    doc.add_paragraph()

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    rr = table2.rows[0].cells
    rr[0].text, rr[1].text = "Responded By:", ""
    rr[2].text, rr[3].text = "Responded Date:", str(rfi.get("response_date", "")) if rfi.get("response_date") else ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _drive_headers():
    token = get_google_token("drive")
    return {"Authorization": f"Bearer {token}"}


def save_rfi_docx_to_drive(rfi: dict, docx_bytes: bytes) -> dict:
    """
    Saves the generated RFI document into the project's real RFIs folder
    (the same folder that already holds that project's RFI Log workbook -
    resolved by locating "RFIs" under a "06 RFI & Submittals"-style parent,
    not a folder ID hardcoded per project, so this works the same way
    across projects with the same real structure).
    """
    headers = _drive_headers()
    project_code = rfi.get("project_code")
    project_root = _project_drive_root(project_code)
    if not project_root:
        return {"error": f"no known Drive root for project {project_code}"}

    rfi_folder_id = _find_rfi_folder(project_root, headers)
    if not rfi_folder_id:
        return {"error": f"no 'RFIs' folder found under project {project_code}'s Drive root"}

    filename = f"{project_code} - {rfi['rfi_number']} - {rfi['subject']}.docx"
    metadata = {"name": filename, "parents": [rfi_folder_id]}
    boundary = "hcirfidoc"
    body = (f"--{boundary}\r\nContent-Type: application/json; charset=UTF-8\r\n\r\n{json.dumps(metadata)}\r\n"
            f"--{boundary}\r\nContent-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n"
            f"Content-Transfer-Encoding: base64\r\n\r\n")
    import base64
    body_bytes = (body.encode("utf-8") + base64.b64encode(docx_bytes) +
                  f"\r\n--{boundary}--".encode("utf-8"))
    r = requests.post("https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart&supportsAllDrives=true",
                       headers={**headers, "Content-Type": f"multipart/related; boundary={boundary}"},
                       data=body_bytes, timeout=60)
    if r.status_code != 200:
        return {"error": f"Drive upload failed: {r.status_code} {r.text[:300]}"}
    return r.json()


# Project Shared Drive IDs (the drive itself, not the 00_Bids subfolder -
# "RFIs"/"06 RFI & Submittals" and other project folders are siblings of
# 00_Bids at the Shared Drive top level, not nested inside it).
_PROJECT_DRIVE_ROOTS = {
    "1355R": "0AAI3pETbQDUUUk9PVA",
    "101F":  "0AEnYRTul6Q7IUk9PVA",
    "64EW":  "0AFf5boDk6qxpUk9PVA",
}


def _project_drive_root(project_code: str) -> str:
    return _PROJECT_DRIVE_ROOTS.get(project_code)


def _find_rfi_folder(project_root: str, headers: dict) -> str:
    """Finds the real 'RFIs' folder under a project's Drive root, wherever
    it actually sits (e.g. 1355R has it nested under "06 RFI & Submittals"),
    by searching the whole project tree rather than assuming a fixed path."""
    r = requests.get("https://www.googleapis.com/drive/v3/files", headers=headers,
                      params={"q": "name = 'RFIs' and mimeType = 'application/vnd.google-apps.folder' and trashed = false",
                              "fields": "files(id,name,parents)", "pageSize": 10,
                              "corpora": "allDrives", "includeItemsFromAllDrives": "true",
                              "supportsAllDrives": "true"}, timeout=30)
    for f in r.json().get("files", []):
        # verify it's actually under this project's tree by walking parents up to the root
        cur = f["id"]
        for _ in range(6):
            rp = requests.get(f"https://www.googleapis.com/drive/v3/files/{cur}", headers=headers,
                               params={"fields": "id,parents", "supportsAllDrives": "true"}, timeout=30)
            j = rp.json()
            parents = j.get("parents") or []
            if not parents:
                break
            if parents[0] == project_root:
                return f["id"]
            cur = parents[0]
    return None


def _find_rfi_log_sheet(project_root: str, headers: dict) -> str:
    rfi_folder = _find_rfi_folder(project_root, headers)
    if not rfi_folder:
        return None
    r = requests.get("https://www.googleapis.com/drive/v3/files", headers=headers,
                      params={"q": f"'{rfi_folder}' in parents and trashed=false and name contains 'RFI' and name contains 'Log'",
                              "fields": "files(id,name)", "supportsAllDrives": "true", "includeItemsFromAllDrives": "true"}, timeout=30)
    files = r.json().get("files", [])
    return files[0]["id"] if files else None


def update_rfi_log_tracker(rfi: dict) -> dict:
    """
    Appends/updates this RFI's row in the project's real RFI Log tab,
    matching its exact existing column schema (RFI Date, Request Number,
    Date Sent, Date Returned, DESCRIPTION, Comments, Response, Sub RFI,
    Who Requested, Last Updated, Party Responsible) rather than inventing a
    new tracker format.

    The real RFI Log workbook is a native .xlsx file sitting in Drive, NOT
    a converted Google Sheet - the Sheets API 400s on it ("document must
    not be an Office file"). Same download/edit-with-openpyxl/re-upload
    pattern used everywhere else this session for native .xlsx trackers,
    not the Sheets API.
    """
    import openpyxl

    project_root = _project_drive_root(rfi.get("project_code"))
    headers = _drive_headers()
    sheet_id = _find_rfi_log_sheet(project_root, headers)
    if not sheet_id:
        return {"error": f"no RFI Log sheet found for project {rfi.get('project_code')}"}

    r = requests.get(f"https://www.googleapis.com/drive/v3/files/{sheet_id}?alt=media&supportsAllDrives=true",
                      headers=headers, timeout=30)
    if r.status_code != 200:
        return {"error": f"could not download RFI Log workbook: {r.status_code}"}

    buf = io.BytesIO(r.content)
    wb = openpyxl.load_workbook(buf)
    log_tab = None
    for name in wb.sheetnames:
        if "log" in name.lower():
            log_tab = name
            break
    if not log_tab:
        return {"error": "no Log tab found in RFI Log workbook"}

    ws = wb[log_tab]
    # find the header row (has "RFI Date" in column A) then the first blank row after it
    header_row = None
    for row_idx in range(1, ws.max_row + 1):
        if str(ws.cell(row=row_idx, column=1).value or "").strip() == "RFI Date":
            header_row = row_idx
            break
    if header_row is None:
        return {"error": "could not find 'RFI Date' header row in Log tab"}

    from openpyxl.cell.cell import MergedCell

    def _row_is_writable(row_idx):
        return all(not isinstance(ws.cell(row=row_idx, column=c), MergedCell) for c in range(1, 12))

    target_row = header_row + 1
    while ws.cell(row=target_row, column=2).value not in (None, "") or not _row_is_writable(target_row):
        target_row += 1
        if target_row > header_row + 500:
            return {"error": "could not find a writable blank row in Log tab within 500 rows"}

    values = [
        str(rfi.get("submitted_date", "")), rfi.get("rfi_number", ""),
        str(rfi.get("submitted_date", "")), str(rfi.get("response_date", "") or ""),
        rfi.get("subject", ""), "", rfi.get("response", "") or "", "",
        rfi.get("submitted_by", ""), datetime.now().strftime("%Y-%m-%d"), rfi.get("submitted_to", ""),
    ]
    for col, val in enumerate(values, start=1):
        ws.cell(row=target_row, column=col).value = val

    out = io.BytesIO()
    wb.save(out)
    out_bytes = out.getvalue()

    r2 = requests.patch(
        f"https://www.googleapis.com/upload/drive/v3/files/{sheet_id}?uploadType=media&supportsAllDrives=true",
        headers={**headers, "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
        data=out_bytes, timeout=60)
    if r2.status_code != 200:
        return {"error": f"upload failed: {r2.status_code} {r2.text[:300]}"}
    return {"ok": True, "row": target_row, "tab": log_tab}


BUCK_EMAIL = "buck@hendricksoninc.com"


def _resolve_recipient_gate(to_email: str | None) -> dict:
    """
    Server-side is_onboarded enforcement. Buck's docstring in gbt_gateway.py
    (GET /users) claims TO-routing already gates on is_onboarded - it didn't;
    to_email was passed straight through from the caller with no check. This
    closes that gap (found + authorized to fix 2026-07-10).

    Scope: only applies when to_email matches a known internal platform_users
    row - i.e. someone Field GPT might address directly as an HCI team
    member. External RFI recipients (architects, subs - not in platform_users
    at all) are unaffected; the RFI-to-external-party path is the normal,
    expected use of this workflow, not something onboarding status governs.
    """
    if not to_email:
        return {"email": BUCK_EMAIL, "redirected": False, "reason": "no_recipient_given"}
    try:
        with _pg() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT actor_name, is_onboarded FROM platform_users WHERE email = %s AND active = TRUE",
                    (to_email,)
                )
                row = cur.fetchone()
    except Exception:
        row = None
    if row and not row["is_onboarded"]:
        return {"email": BUCK_EMAIL, "redirected": True,
                "reason": f"{row['actor_name']} is a known team member but not yet onboarded"}
    return {"email": to_email, "redirected": False, "reason": None}


def create_rfi_email_draft(rfi: dict, docx_bytes: bytes, to_email: str, to_name: str) -> dict:
    """Creates an Outlook draft (never auto-sent) with the generated RFI
    document attached. Buck's explicit routing correction, 2026-07-11: every
    draft goes To: Buck, CC: the real intended recipient (when known) - not
    To: the recipient with Buck BCC'd, and not To: Buck alone with the
    recipient buried only in body prose. He reviews the draft, sees exactly
    who it's headed to via the CC line, and sends it on himself. This
    replaces the BCC approach from 2026-07-10 - that one didn't give Buck
    structural visibility into who a "defaulted" draft was actually for."""
    from microsoft_graph import create_draft, add_attachment_bytes

    subject = f"RFI {rfi['rfi_number']} - {rfi['subject']} - {rfi.get('project_name', '')}"
    greet_name = to_name if (to_email and to_email.lower() != BUCK_EMAIL) else "Buck"
    body_html = (
        f"<p>Hi {greet_name},</p>"
        f"<p>Please see the attached Request for Information ({rfi['rfi_number']}) "
        f"for {rfi.get('project_name', '')}.</p>"
        f"<p>{rfi.get('question', '')}</p>"
        f"<p>Please respond by replying to this email or returning the attached form.</p>"
        f"<p>Thanks,<br>Buck Adams<br>Hendrickson Construction</p>"
    )
    cc = None if (not to_email or to_email.lower() == BUCK_EMAIL) else [(to_name, to_email)]
    draft = create_draft(subject, body_html, [("Buck Adams", BUCK_EMAIL)], cc=cc)
    draft_id = draft["id"]
    filename = f"{rfi['rfi_number']} - {rfi['subject']}.docx"
    _, err = add_attachment_bytes(draft_id, filename,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", docx_bytes)
    if err:
        return {"draft": draft, "attachment_error": err}
    return {"draft": draft, "attached": filename}


def run_rfi_workflow(rfi_id: int, to_email: str = None, to_name: str = None) -> dict:
    """
    Full pipeline: read RFI -> generate Word doc -> save to Drive -> update
    Log tracker -> create Outlook draft with attachment. Returns evidence
    for every step (per GBT's evidence-first acceptance requirement) rather
    than a single success/fail flag.
    """
    evidence = {"rfi_id": rfi_id, "steps": {}}
    rfi = get_rfi(rfi_id)
    if not rfi:
        return {"error": f"RFI {rfi_id} not found"}
    evidence["steps"]["read_rfi"] = {"ok": True, "rfi_number": rfi["rfi_number"], "subject": rfi["subject"]}

    docx_bytes = generate_rfi_docx(rfi)
    evidence["steps"]["generate_docx"] = {"ok": True, "size_bytes": len(docx_bytes)}

    drive_result = save_rfi_docx_to_drive(rfi, docx_bytes)
    evidence["steps"]["save_to_drive"] = (
        {"ok": True, "file_id": drive_result.get("id"), "file_name": drive_result.get("name")}
        if "id" in drive_result else {"ok": False, "error": drive_result.get("error")}
    )

    tracker_result = update_rfi_log_tracker(rfi)
    evidence["steps"]["update_tracker"] = (
        {"ok": True} if "error" not in tracker_result else {"ok": False, "error": tracker_result["error"]}
    )

    # Buck, 2026-07-10: "I am the only email. I will read and swap out the
    # email and copy who it needs to go to." Buck, 2026-07-11 (correcting the
    # first version of this): "the draft goes to me and I put the person it
    # is directed and gets cc'd to in my draft" - every draft goes To: Buck,
    # CC: the real intended recipient when one is known (create_rfi_email_draft
    # handles this), not the To:recipient/BCC:Buck shape this used to build.
    # is_onboarded gate still runs and is recorded as evidence - it doesn't
    # change the CC-routing behavior below since Buck is the only onboarded
    # user right now, but keeps the audit trail of who's a known-not-yet-
    # onboarded team member vs. a genuine external party.
    gate = _resolve_recipient_gate(to_email)
    defaulted_to_buck = not to_email or gate["redirected"]

    email_result = create_rfi_email_draft(rfi, docx_bytes, to_email, to_name or to_email)
    evidence["steps"]["create_email_draft"] = (
        {"ok": True, "draft_id": email_result["draft"]["id"], "attached": email_result.get("attached"),
         "defaulted_to_buck": defaulted_to_buck, "redirect_reason": gate["reason"],
         "cc": to_email if (to_email and to_email.lower() != BUCK_EMAIL) else None}
        if "draft" in email_result else {"ok": False, "error": str(email_result)}
    )

    evidence["all_steps_ok"] = all(s.get("ok") for s in evidence["steps"].values())

    # Persist the generated document + email draft as durable links on the RFI
    # row itself, not just one-time API response evidence - per GBT's "tracker
    # becomes the master record" architecture. Best-effort: a failure here
    # doesn't roll back the real Drive/email actions already taken above.
    doc_id = evidence["steps"]["save_to_drive"].get("file_id")
    draft_id = evidence["steps"]["create_email_draft"].get("draft_id")
    if doc_id or draft_id:
        try:
            update_rfi(rfi_id, updated_by="rfi_workflow",
                       rfi_document_id=doc_id, email_draft_id=draft_id)
            evidence["steps"]["persist_links"] = {"ok": True, "rfi_document_id": doc_id, "email_draft_id": draft_id}
        except Exception as e:
            evidence["steps"]["persist_links"] = {"ok": False, "error": str(e)}

    return evidence
